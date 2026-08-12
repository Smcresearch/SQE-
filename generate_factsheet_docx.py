"""
Generate SQE_Factsheet.docx — Word version of the SQE (SMC Quant Equity) factsheet.

Reads the same data.js source as generate_factsheet.py so the two documents can
never drift apart, and lays the content out with python-docx tables (Word has no
CSS, so every "card" / bar / grid in the HTML becomes a shaded table here).
"""
import json, tempfile
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(r'd:/SQE-host')
DOCX_OUT = ROOT / 'SQE_Factsheet.docx'

# ── Palette (mirrors the CSS custom properties in generate_factsheet.py) ──
PRI   = '0F2B54'
PRI2  = '1565C0'
ACC   = '16C784'
RED   = 'EA3943'
BG    = 'F3F6FB'
BDR   = 'DDE3EF'
TXT   = '1A1A2E'
SUB   = '6B7A99'
WHITE = 'FFFFFF'

FONT   = 'Calibri'
MONO   = 'Consolas'          # also the only one of the two carrying U+2588 for the bars
SYMBOL = 'Segoe UI Symbol'


# ══════════════════════════════ data ══════════════════════════════
with open(ROOT / 'data.js', 'r', encoding='utf-8') as f:
    content = f.read()

clean = content.strip()
if clean.startswith('const DASHBOARD_DATA ='):
    clean = clean[len('const DASHBOARD_DATA ='):].strip()
if clean.endswith(';'):
    clean = clean[:-1].strip()
data = json.loads(clean)

portfolio    = data['total759']['current_portfolio']
base_metrics = data['total759']['layer_metrics']['Base']
exec_sum     = data['total759']['exec_summary']

last_update = data.get('last_update', 'N/A')
try:
    last_update_fmt = datetime.strptime(last_update[:10], '%Y-%m-%d').strftime('%B %d, %Y')
except ValueError:
    last_update_fmt = last_update

M = {
    'CAGR': base_metrics['CAGR'],
    'Bench_CAGR': exec_sum['CAGR']['Bench'] * 100,
    'Bench_N50_CAGR': data['nifty50']['exec_summary']['CAGR']['Bench'] * 100,
    'Alpha': exec_sum['Alpha vs Bench']['Base'] * 100,
    'Volatility': base_metrics['Volatility'],
    'Sharpe': base_metrics['Sharpe'],
    'Sortino': base_metrics['Sortino'],
    'Calmar': base_metrics['Calmar'],
    'Max_DD': base_metrics['Max_DD'],
    'Win_Rate': base_metrics['Win_Rate'],
    'Avg_Gain': base_metrics['Avg_Gain'],
    'Avg_Loss': base_metrics['Avg_Loss'],
    'Total_Return': base_metrics['Total_Return'],
    'Best_Month': exec_sum['Best Month']['Base'] * 100,
    'Worst_Month': exec_sum['Worst Month']['Base'] * 100,
    'VaR_95': exec_sum['VaR 95%']['Base'] * 100,
}

sector_map = {}
for h in portfolio:
    sector_map[h['sector']] = sector_map.get(h['sector'], 0) + h['weight'] * 100
sector_map = dict(sorted(sector_map.items(), key=lambda x: -x[1]))

SECTOR_COLORS = ['1A73E8', '34A853', 'FBBC04', 'EA4335', '9C27B0',
                 '00BCD4', 'FF5722', '607D8B', '795548', 'E91E63']

top10_weight = sum(h['weight'] for h in sorted(portfolio, key=lambda x: -x['weight'])[:10]) * 100


# ══════════════════════════ docx helpers ══════════════════════════
# WordprocessingML validates child order strictly — Word refuses to open the file if a
# property element lands out of sequence, so every insert below names its successors.
_TC_AFTER_SHD = ('w:noWrap', 'w:tcMar', 'w:textDirection', 'w:tcFitText', 'w:vAlign',
                 'w:hideMark', 'w:cellIns', 'w:cellDel', 'w:cellMerge')
_TC_AFTER_MAR = ('w:textDirection', 'w:tcFitText', 'w:vAlign', 'w:hideMark',
                 'w:cellIns', 'w:cellDel', 'w:cellMerge')


def shade(cell, hex_fill):
    tcPr = cell._tc.get_or_add_tcPr()
    for old in tcPr.findall(qn('w:shd')):
        tcPr.remove(old)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_fill)
    tcPr.insert_element_before(shd, *_TC_AFTER_SHD)


def cell_margins(cell, top=60, bottom=60, left=110, right=110):
    """Cell padding, in twentieths of a point."""
    tcPr = cell._tc.get_or_add_tcPr()
    for old in tcPr.findall(qn('w:tcMar')):
        tcPr.remove(old)
    tcMar = OxmlElement('w:tcMar')
    # w:left/w:right, not w:start/w:end — the latter are ECMA-376 2nd edition and
    # Word 2007 refuses to open a document that uses them.
    for tag, val in (('top', top), ('left', left), ('bottom', bottom), ('right', right)):
        el = OxmlElement(f'w:{tag}')
        el.set(qn('w:w'), str(val))
        el.set(qn('w:type'), 'dxa')
        tcMar.append(el)
    tcPr.insert_element_before(tcMar, *_TC_AFTER_MAR)


def table_borders(table, color=BDR, size=4, edges=('top', 'left', 'bottom', 'right',
                                                   'insideH', 'insideV')):
    borders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = OxmlElement(f'w:{edge}')
        if edge in edges:
            el.set(qn('w:val'), 'single')
            el.set(qn('w:sz'), str(size))
            el.set(qn('w:color'), color)
        else:
            el.set(qn('w:val'), 'none')
            el.set(qn('w:sz'), '0')
        el.set(qn('w:space'), '0')
        borders.append(el)
    tblPr = table._tbl.tblPr
    for old in tblPr.findall(qn('w:tblBorders')):
        tblPr.remove(old)
    tblPr.insert_element_before(borders, 'w:shd', 'w:tblLayout', 'w:tblCellMar',
                                'w:tblLook', 'w:tblCaption', 'w:tblDescription',
                                'w:tblPrChange')


def no_borders(table):
    table_borders(table, edges=())


def write(cell_or_par, text, *, size=10, bold=False, color=TXT, font=FONT,
          align=None, space_after=0, space_before=0, caps=False, italic=False,
          append=False):
    """Write a run into a cell (replacing its empty first paragraph) or a paragraph."""
    if hasattr(cell_or_par, 'paragraphs'):          # a table cell
        par = cell_or_par.paragraphs[0] if (not append or len(cell_or_par.paragraphs) == 1
                                            and not cell_or_par.paragraphs[0].runs) \
            else cell_or_par.add_paragraph()
    else:
        par = cell_or_par
    run = par.add_run(text.upper() if caps else text)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.name = font
    run.font.color.rgb = RGBColor.from_string(color)
    # East-Asian font mapping, else Word may substitute for the mono runs.
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font)
    if align is not None:
        par.alignment = align
    par.paragraph_format.space_after = Pt(space_after)
    par.paragraph_format.space_before = Pt(space_before)
    return par


def add_table(doc, rows, cols, widths=None):
    t = doc.add_table(rows=rows, cols=cols)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = w
    return t


def section_title(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text.upper())
    r.font.size = Pt(10)
    r.font.bold = True
    r.font.name = FONT
    r.font.color.rgb = RGBColor.from_string(PRI)
    r.font.all_caps = True
    # Rule under the heading, standing in for the CSS ::after flex line.
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '2')
    bottom.set(qn('w:color'), BDR)
    pbdr.append(bottom)
    pPr.insert_element_before(pbdr, 'w:shd', 'w:tabs', 'w:suppressAutoHyphens', 'w:kinsoku',
                              'w:wordWrap', 'w:overflowPunct', 'w:topLinePunct',
                              'w:autoSpaceDE', 'w:autoSpaceDN', 'w:bidi', 'w:adjustRightInd',
                              'w:snapToGrid', 'w:spacing', 'w:ind', 'w:contextualSpacing',
                              'w:mirrorIndents', 'w:suppressOverlap', 'w:jc',
                              'w:textDirection', 'w:textAlignment', 'w:textboxTightWrap',
                              'w:outlineLvl', 'w:divId', 'w:cnfStyle', 'w:rPr', 'w:sectPr',
                              'w:pPrChange')
    return p


def spacer(doc, pts=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    p.add_run('').font.size = Pt(pts)
    return p


def note_box(doc, blocks, fill='F9FAFC', size=8.5, color=SUB):
    """Grey disclaimer-style box. `blocks` is a list of [(text, bold, color), ...] lines."""
    t = add_table(doc, 1, 1, [Cm(18.4)])
    table_borders(t, color=BDR)
    c = t.cell(0, 0)
    shade(c, fill)
    cell_margins(c, top=120, bottom=120, left=180, right=180)
    first = True
    for block in blocks:
        par = c.paragraphs[0] if first else c.add_paragraph()
        first = False
        par.paragraph_format.space_after = Pt(6)
        par.paragraph_format.space_before = Pt(0)
        for text, bold, col in block:
            r = par.add_run(text)
            r.font.size = Pt(size)
            r.font.bold = bold
            # Calibri has no U+26A0; name the symbol font so Word does not draw a box.
            r.font.name = SYMBOL if '⚠' in text else FONT
            r.font.color.rgb = RGBColor.from_string(col or color)
    return t


def page_break(doc):
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


# ══════════════════════════ document ══════════════════════════
doc = Document()

st = doc.styles['Normal']
st.font.name = FONT
st.font.size = Pt(10)
st.element.rPr.rFonts.set(qn('w:eastAsia'), FONT)
st.paragraph_format.space_after = Pt(4)
st.paragraph_format.line_spacing = 1.15

sec = doc.sections[0]
sec.top_margin = Cm(1.2)
sec.bottom_margin = Cm(1.2)
sec.left_margin = Cm(1.3)
sec.right_margin = Cm(1.3)

W = Cm(18.4)   # usable content width

# ── Logo (Word cannot embed webp; flatten onto white and convert to PNG) ──
logo_src = ROOT / 'smc_logo.webp'
if logo_src.exists():
    try:
        from PIL import Image
        img = Image.open(logo_src).convert('RGBA')
        flat = Image.new('RGBA', img.size, (255, 255, 255, 255))
        flat.alpha_composite(img)
        with tempfile.TemporaryDirectory() as tmp:
            logo_png = Path(tmp) / 'smc_logo.png'
            flat.convert('RGB').save(logo_png)
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(8)
            p.add_run().add_picture(str(logo_png), height=Cm(1.15))
    except Exception as e:      # noqa: BLE001 - logo is decorative, never fatal
        print(f'WARNING: logo skipped ({e})')

# ── Blue header banner ──
hdr = add_table(doc, 6, 4, [Cm(4.6)] * 4)
no_borders(hdr)
for row in hdr.rows:
    for c in row.cells:
        shade(c, PRI)
        cell_margins(c, top=40, bottom=40, left=80, right=80)

for r_i in (0, 1, 2, 5):
    hdr.cell(r_i, 0).merge(hdr.cell(r_i, 3))

write(hdr.cell(0, 0), 'F A C T S H E E T', size=8, bold=True, color='9EC3EE',
      align=WD_ALIGN_PARAGRAPH.CENTER, space_before=6)
write(hdr.cell(1, 0), 'SQE SmallCase Terminal', size=22, bold=True, color=WHITE,
      align=WD_ALIGN_PARAGRAPH.CENTER)
write(hdr.cell(2, 0), 'SMC Quant Equity — All Indices Portfolio', size=11, color='C7DCF5',
      align=WD_ALIGN_PARAGRAPH.CENTER, space_after=8)

pills = [(f"{M['CAGR']:.2f}%", 'CAGR', '4ADE80'),
         (f"{M['Total_Return']:.1f}%", 'Total Return', '4ADE80'),
         ('High Volatility', 'Risk Level', WHITE),
         ('Long Term', 'Horizon', WHITE)]
for i, (val, lbl, col) in enumerate(pills):
    write(hdr.cell(3, i), val, size=15 if col == '4ADE80' else 11, bold=True, color=col,
          align=WD_ALIGN_PARAGRAPH.CENTER)
    write(hdr.cell(4, i), lbl, size=8, color='A9C6E8', align=WD_ALIGN_PARAGRAPH.CENTER)

write(hdr.cell(5, 0),
      f'Last updated: {last_update_fmt}   ·   smcresearch.github.io/SQE-/',
      size=8, color='9EC3EE', align=WD_ALIGN_PARAGRAPH.CENTER, space_before=6, space_after=6)

# ── Portfolio overview ──
section_title(doc, 'Portfolio Overview')
overview = [
    ('Portfolio Type', 'Thematic / Quant', TXT),
    ('Constituents', 'Indian Stocks (NSE)', TXT),
    ('Asset Class', 'Equity Multi Cap', TXT),
    ('Universe', 'All NSE Indices', TXT),
    ('No. of Stocks', str(len(portfolio)), TXT),
    ('Launch Period', 'January 2020', TXT),
    ('CAGR (Portfolio)', f"{M['CAGR']:.2f}%", ACC),
    ('CAGR (Nifty 500)', f"{M['Bench_CAGR']:.2f}%", PRI2),
]
ov = add_table(doc, 2, 4, [Cm(4.6)] * 4)
table_borders(ov, color=BDR)
for i, (k, v, col) in enumerate(overview):
    c = ov.cell(i // 4, i % 4)
    shade(c, BG)
    cell_margins(c)
    write(c, k.upper(), size=7.5, bold=True, color=SUB)
    write(c, v, size=11, bold=True, color=col, append=True)

# ── Rationale ──
section_title(doc, 'Portfolio Rationale')
rat = add_table(doc, 1, 1, [W])
table_borders(rat, color='C8D8F0')
rc = rat.cell(0, 0)
shade(rc, 'F0F5FF')
cell_margins(rc, top=140, bottom=140, left=200, right=200)

p = rc.paragraphs[0]
p.paragraph_format.space_after = Pt(6)
for text, bold in [('SQE (SMC Quant Equity)', True),
                   (' is a concentrated, research-backed equity portfolio designed for ', False),
                   ('long-term wealth creation', True),
                   ('. It aims to deliver superior risk-adjusted returns by investing in a select '
                    'basket of high-quality Indian stocks across all major NSE indices.', False)]:
    r = p.add_run(text)
    r.font.size = Pt(10)
    r.font.bold = bold
    r.font.name = FONT
    r.font.color.rgb = RGBColor.from_string(PRI if bold else TXT)

bullets = [
    [('Stocks are selected using a ', False),
     ('proprietary quantitative model', True),
     (' developed by SMC Research that evaluates each stock\u2019s risk-return profile relative '
      'to the broader market', False)],
    [('The portfolio holds ', False),
     (f'{len(portfolio)} high-conviction positions', True),
     (' drawn from the full universe of NSE-listed companies — large, mid and small cap — with '
      'individual position sizes capped at ', False),
     ('10%', True), (' to control single-stock risk', False)],
    [('Every month, the portfolio is ', False),
     ('systematically rebalanced', True),
     (' to capture new opportunities and manage risk — there is no discretionary or emotional '
      'decision-making', False)],
    [('The model has consistently generated ', False),
     ('alpha over the Nifty 500 benchmark', True),
     (' since inception, with a disciplined focus on both upside capture and downside protection',
      False)],
]
for b in bullets:
    par = rc.add_paragraph(style='List Bullet')
    par.paragraph_format.space_after = Pt(3)
    par.paragraph_format.left_indent = Cm(0.7)
    for text, bold in b:
        r = par.add_run(text)
        r.font.size = Pt(9.5)
        r.font.bold = bold
        r.font.name = FONT
        r.font.color.rgb = RGBColor.from_string(PRI if bold else TXT)

# ── Rebalance schedule ──
section_title(doc, 'Rebalance Schedule')
reb = [('Frequency', 'Monthly'), ('Rebalance Day', '1st Trading Day'),
       ('Last Rebalance', 'August 2026'), ('Next Rebalance', 'September 2026'),
       ('Managed By', 'SMC Research')]
rt = add_table(doc, 1, 5, [Cm(3.68)] * 5)
table_borders(rt, color=BDR)
for i, (k, v) in enumerate(reb):
    c = rt.cell(0, i)
    shade(c, BG)
    cell_margins(c)
    write(c, k.upper(), size=7.5, bold=True, color=SUB, align=WD_ALIGN_PARAGRAPH.CENTER)
    write(c, v, size=10, bold=True, color=PRI, align=WD_ALIGN_PARAGRAPH.CENTER, append=True)


# ═══════════════════ PAGE 2 ═══════════════════
page_break(doc)

section_title(doc, 'Performance & Risk Metrics')
left_metrics = [
    ('CAGR (Portfolio)', f"{M['CAGR']:.2f}%", ACC),
    ('CAGR (Nifty 500 — Benchmark)', f"{M['Bench_CAGR']:.2f}%", TXT),
    ('CAGR (Nifty 50 — Reference)', f"{M['Bench_N50_CAGR']:.2f}%", TXT),
    ('Total Return', f"{M['Total_Return']:.2f}%", ACC),
    ('Alpha vs Nifty 500', f"+{M['Alpha']:.2f}%", ACC),
    ('Annualised Volatility', f"{M['Volatility']:.2f}%", TXT),
    ('Sharpe Ratio', f"{M['Sharpe']:.2f}", ACC),
    ('Sortino Ratio', f"{M['Sortino']:.2f}", ACC),
    ('Calmar Ratio', f"{M['Calmar']:.2f}", ACC),
]
right_metrics = [
    ('Max Drawdown', f"{M['Max_DD']:.2f}%", RED),
    ('Win Rate (Monthly)', f"{M['Win_Rate']:.1f}%", ACC),
    ('Avg Monthly Gain', f"+{M['Avg_Gain']:.2f}%", ACC),
    ('Avg Monthly Loss', f"{M['Avg_Loss']:.2f}%", RED),
    ('Best Month', f"+{M['Best_Month']:.2f}%", ACC),
    ('Worst Month', f"{M['Worst_Month']:.2f}%", RED),
    ('VaR (95%, Monthly)', f"{M['VaR_95']:.2f}%", RED),
    ('Live Since', 'Jan 2020', TXT),
    ('', '', TXT),
]
mt = add_table(doc, len(left_metrics), 4, [Cm(5.6), Cm(3.6), Cm(5.6), Cm(3.6)])
table_borders(mt, color=BDR, edges=('insideH',))
for i, (l, r) in enumerate(zip(left_metrics, right_metrics)):
    for off, (k, v, col) in ((0, l), (2, r)):
        cell_margins(mt.cell(i, off), top=50, bottom=50, left=40, right=40)
        cell_margins(mt.cell(i, off + 1), top=50, bottom=50, left=40, right=40)
        write(mt.cell(i, off), k, size=9.5, color=SUB)
        write(mt.cell(i, off + 1), v, size=10, bold=True, color=col, font=MONO,
              align=WD_ALIGN_PARAGRAPH.RIGHT)

# ── Sector allocation (CSS bars become block-character runs) ──
section_title(doc, 'Sector Allocation')
max_wt = max(sector_map.values())
at = add_table(doc, len(sector_map), 3, [Cm(6.2), Cm(10.0), Cm(2.2)])
no_borders(at)
for i, (sector, wt) in enumerate(sector_map.items()):
    col = SECTOR_COLORS[i % len(SECTOR_COLORS)]
    cell_margins(at.cell(i, 0), top=30, bottom=30, left=40, right=40)
    cell_margins(at.cell(i, 1), top=30, bottom=30, left=40, right=40)
    cell_margins(at.cell(i, 2), top=30, bottom=30, left=40, right=40)
    write(at.cell(i, 0), sector, size=9, color=TXT)
    bar_len = max(1, round(wt / max_wt * 38))
    par = at.cell(i, 1).paragraphs[0]
    par.paragraph_format.space_after = Pt(0)
    r = par.add_run('█' * bar_len)
    r.font.size = Pt(8)
    r.font.name = MONO
    r.font.color.rgb = RGBColor.from_string(col)
    r._element.rPr.rFonts.set(qn('w:eastAsia'), MONO)
    write(at.cell(i, 2), f'{wt:.1f}%', size=9, bold=True, font=MONO,
          align=WD_ALIGN_PARAGRAPH.RIGHT)


# ═══════════════════ PAGE 3 ═══════════════════
page_break(doc)

section_title(doc, f'Current Holdings — {len(portfolio)} Stocks · August 2026')
cols = [Cm(1.0), Cm(5.0), Cm(6.0), Cm(2.1), Cm(2.5), Cm(1.8)]
ht = add_table(doc, len(portfolio) + 1, 6, cols)
table_borders(ht, color=BDR, edges=('insideH', 'bottom'))
headers = ['#', 'STOCK', 'SECTOR', 'WEIGHT', 'LTP (₹)', 'DAY CHG']
for i, h in enumerate(headers):
    c = ht.cell(0, i)
    shade(c, PRI)
    cell_margins(c, top=70, bottom=70, left=70, right=70)
    write(c, h, size=8, bold=True, color=WHITE,
          align=WD_ALIGN_PARAGRAPH.CENTER if i in (0, 3, 4, 5) else None)
ht.rows[0]._tr.get_or_add_trPr().append(OxmlElement('w:tblHeader'))   # repeat on page 2

for idx, h in enumerate(portfolio, 1):
    row = ht.rows[idx]
    chg = h.get('change_pct', 0)
    chg_col = ACC if chg >= 0 else RED
    chg_s = f'+{chg:.2f}%' if chg >= 0 else f'{chg:.2f}%'
    name = h['clean_symbol'] + (' [NEW]' if h.get('status') == 'Added' else '')
    vals = [
        (str(idx), 8, False, SUB, FONT, WD_ALIGN_PARAGRAPH.CENTER),
        (name, 9.5, True, PRI, FONT, None),
        (h['sector'], 8.5, False, TXT, FONT, None),
        (f"{h['weight'] * 100:.1f}%", 9, True, TXT, MONO, WD_ALIGN_PARAGRAPH.CENTER),
        (f"₹{h.get('ltp', 0):,.2f}", 9, False, TXT, MONO, WD_ALIGN_PARAGRAPH.RIGHT),
        (chg_s, 9, False, chg_col, MONO, WD_ALIGN_PARAGRAPH.RIGHT),
    ]
    for i, (text, size, bold, col, font, align) in enumerate(vals):
        c = row.cells[i]
        cell_margins(c, top=45, bottom=45, left=70, right=70)
        if idx % 2 == 0:
            shade(c, 'FAFBFD')
        write(c, text, size=size, bold=bold, color=col, font=font, align=align)


# ═══════════════════ PAGE 4 ═══════════════════
page_break(doc)

section_title(doc, 'How It Works')
steps = [
    ('Step 1', 'SMC Research runs the proprietary quant model every month'),
    ('Step 2', 'Updated portfolio with buy/sell actions published on the dashboard'),
    ('Step 3', 'Execute the rebalance trades at the start of each month'),
    ('Step 4', 'Hold for 3–5 years for best risk-adjusted returns'),
]
stt = add_table(doc, 1, 4, [Cm(4.6)] * 4)
table_borders(stt, color=BDR)
for i, (k, v) in enumerate(steps):
    c = stt.cell(0, i)
    shade(c, BG)
    cell_margins(c, top=110, bottom=110)
    write(c, k.upper(), size=8, bold=True, color=PRI2, align=WD_ALIGN_PARAGRAPH.CENTER)
    write(c, v, size=9, color=SUB, align=WD_ALIGN_PARAGRAPH.CENTER, append=True)

section_title(doc, 'Key Risk Factors')
risks = [
    ('Market Risk',
     'Equity prices can decline sharply due to macroeconomic, geopolitical or company-specific '
     f"factors. The portfolio has seen a maximum drawdown of {M['Max_DD']:.2f}% over the "
     'backtest period.'),
    ('Concentration Risk',
     f'The portfolio holds {len(portfolio)} stocks, but is top-heavy — the largest positions are '
     f'capped at 10% each and the top 10 holdings account for roughly {top10_weight:.0f}% of the '
     'portfolio. Sector concentration may arise during certain market phases.'),
    ('Model Risk',
     'Past performance is not a guarantee of future results. Quantitative models may underperform '
     'during regime changes or unprecedented market events.'),
    ('Liquidity Risk',
     'All constituents are NSE-listed. As the universe spans mid and small cap companies alongside '
     'large caps, some holdings may trade with lower volumes, and liquidity can be temporarily '
     'limited during extreme market stress.'),
]
rkt = add_table(doc, 2, 2, [Cm(9.2)] * 2)
table_borders(rkt, color=BDR)
for i, (k, v) in enumerate(risks):
    c = rkt.cell(i // 2, i % 2)
    shade(c, BG)
    cell_margins(c, top=110, bottom=110, left=180, right=180)
    write(c, k.upper(), size=8, bold=True, color=RED)
    write(c, v, size=9, color=SUB, append=True)

section_title(doc, 'Definitions and Disclosures')
definitions = [
    ('CAGR', 'Compound Annual Growth Rate is a measure of the growth of a portfolio. Returns '
             'generated each year differ; CAGR expresses them as the single annual rate that '
             'would produce the same terminal value over the period. For example, a portfolio '
             'returning 5%, 15% and −7% over three years has a CAGR of 3.94%. In this factsheet '
             'CAGR is computed on backtested model data from January 2020 onwards.'),
    ('Volatility Label', 'Daily changes in stock prices cause fluctuation in the value of your '
             'investment. Each portfolio is categorised into one of three buckets — High, Medium '
             'or Low Volatility — by comparing the portfolio\u2019s volatility against that of the '
             f"Nifty 100 Index. This portfolio\u2019s annualised volatility is {M['Volatility']:.2f}%, "
             'placing it in the High Volatility bucket. High Volatility means changes in your '
             'investment value can be sudden and significant.'),
    ('Investment Horizon', 'The manager\u2019s recommended holding duration. Short Term: <1 year. '
             'Medium Term: 1–3 years. Long Term: >3 years. This portfolio is recommended as '
             'Long Term.'),
    ('Asset Class', 'Constituents are selected from a universe defined by the manager, and that '
             'universe is labelled the Asset Class. All NSE-listed stocks are ranked in decreasing '
             'order of market capitalisation: ranks 1–100 are Large Cap, 101–250 are Mid Cap, and '
             'above 250 are Small Cap. This portfolio is Equity Multi Cap, meaning constituents '
             'may be drawn from more than two of the Large, Mid and Small Cap categories.'),
    ('Rebalance', 'The process of periodically reviewing and updating the constituents of a '
             'portfolio, so that the holdings continue to reflect the underlying strategy. This '
             'portfolio is rebalanced monthly, on the first trading day of each month.'),
    ('Holdings Distribution', 'Constituents are grouped into segments, and the weight of a segment '
             'is the sum of the weights of all constituents in it. For example, if four '
             'constituents of 10% each are Large Cap, the Large Cap segment weight is 40%.'),
    ('Benchmark', 'Portfolio performance in this factsheet is compared against the Nifty 500 '
             f"(CAGR {M['Bench_CAGR']:.2f}% over the same period), which is the index CASE Platforms "
             'designates for the Equity Multi Cap asset class and is therefore the appropriate '
             'comparison for this portfolio. The Nifty 50 (CAGR '
             f"{M['Bench_N50_CAGR']:.2f}%) is shown alongside as a secondary broad-market reference "
             'only. Alpha is stated against the Nifty 500.'),
]
for term, body in definitions:
    par = doc.add_paragraph()
    par.paragraph_format.space_after = Pt(5)
    r = par.add_run(f'{term} — ')
    r.font.size = Pt(9)
    r.font.bold = True
    r.font.name = FONT
    r.font.color.rgb = RGBColor.from_string(PRI)
    r2 = par.add_run(body)
    r2.font.size = Pt(9)
    r2.font.name = FONT
    r2.font.color.rgb = RGBColor.from_string(TXT)

glossary = [
    ('CAGR:', ' Compound Annual Growth Rate — annualised return over the full period.'),
    ('Sharpe Ratio:', ' Risk-adjusted return per unit of total volatility.'),
    ('Sortino Ratio:', ' Risk-adjusted return penalising only downside deviation.'),
    ('Calmar Ratio:', ' CAGR divided by maximum drawdown.'),
    ('Max Drawdown:', ' Largest peak-to-trough decline during the period.'),
    ('Win Rate:', ' Percentage of months with positive returns.'),
    ('Alpha:', ' Portfolio CAGR minus Nifty 500 benchmark CAGR.'),
    ('VaR (95%):', ' Worst expected monthly loss at 95% confidence.'),
]
gt = add_table(doc, 4, 2, [Cm(9.2)] * 2)
no_borders(gt)
for i, (term, body) in enumerate(glossary):
    c = gt.cell(i // 2, i % 2)
    cell_margins(c, top=30, bottom=30, left=0, right=120)
    par = c.paragraphs[0]
    par.paragraph_format.space_after = Pt(0)
    r = par.add_run(term)
    r.font.size = Pt(9)
    r.font.bold = True
    r.font.name = FONT
    r.font.color.rgb = RGBColor.from_string(TXT)
    r2 = par.add_run(body)
    r2.font.size = Pt(9)
    r2.font.name = FONT
    r2.font.color.rgb = RGBColor.from_string(TXT)

section_title(doc, 'General Investment Disclosure')
note_box(doc, [
    [('⚠ IMPORTANT: ', True, RED),
     ('This factsheet is for ', False, SUB), ('informational purposes only', True, SUB),
     (' and does not constitute investment advice, solicitation, or a recommendation to buy or '
      'sell any securities. Past performance is not indicative of future results. Investments in '
      'the securities market are subject to market risks. Read all related documents carefully '
      'before investing.', False, SUB)],
    [('The performance data shown is based on a quantitative model simulation. Live performance '
      'may differ from model results due to transaction costs, taxes (STT, GST), impact cost, and '
      'execution timing. The portfolio is rebalanced monthly at month-open prices. ', False, SUB),
     ('All returns, CAGR and risk figures shown in this factsheet are derived from backtested '
      'model data covering January 2020 onwards, and do not represent the returns of an actual '
      'live-traded portfolio.', True, SUB),
     (' Backtested results are hypothetical, are computed with the benefit of hindsight, and have '
      'inherent limitations. They have not been validated by an independent chartered accountant, '
      'nor verified by the Past Risk and Return Verification Agency (PaRRVA) or any other agency '
      'recognised by SEBI.', False, SUB)],
    [('The volatility label (High/Medium/Low) is determined by comparing the portfolio\u2019s daily '
      'volatility against the Nifty 100 Index. High Volatility means that changes in your '
      'investment value can be sudden and significant.', False, SUB)],
])

section_title(doc, 'Risk Disclosure')
note_box(doc, [
    [('Investing in securities involves various types of risk that may impact your investment. Key '
      'risks affecting all asset classes include changes in: market volatility; general market '
      'conditions; trading volumes, liquidity and settlement periods; interest rates; the rate of '
      'inflation; domestic and global political, economic and financial developments; and '
      'policies, legal or regulatory frameworks set by government and other appropriate '
      'authorities.', False, SUB)],
    [('Risks relating to equity and equity-linked investments:', True, SUB),
     (' equity shares and equity-related instruments are volatile and prone to price fluctuation '
      'on a daily basis. Prices may be affected by trading volume volatility, currency exchange '
      'rates, company specific news and rumours, and other factors. ', False, SUB),
     ('Mid cap and small cap stocks generally exhibit higher volatility than large cap stocks.',
      True, SUB),
     (' As this is a Multi Cap portfolio that draws from the full NSE universe, a meaningful '
      'portion of the holdings may fall in the mid and small cap segments at any given time.',
      False, SUB)],
    [('In light of the risks involved, you should transact in securities only after understanding '
      'the associated risks. Please consider and assess all risk factors and your own risk '
      'tolerance before making investment decisions.', False, SUB)],
])

section_title(doc, 'Manager Disclosure')
note_box(doc, [
    [('SMC Global Securities Ltd.', True, SUB),
     (' is registered with SEBI as a Research Analyst, with its registered office at 11/6B, Shanti '
      'Chamber, Pusa Road, New Delhi – 110005. Registration granted by SEBI and certification from '
      'NISM in no way guarantee performance of the intermediary or provide any assurance of '
      'returns to investors.', False, SUB)],
    [('The content and data available in this material, including index values, return numbers and '
      'rationale, are for information and illustration purposes only. Charts and performance '
      'numbers do not include the impact of transaction fees and other related costs. Past '
      'performance does not guarantee future returns and the performance of the portfolio is '
      'subject to market risk. Data used for the calculation of historical returns and other '
      'information is sourced from exchange-approved third party vendors and has neither been '
      'audited nor independently validated.', False, SUB)],
    [('Information presented in this material shall not be considered a recommendation or '
      'solicitation of an investment. Investors are responsible for their own investment decisions '
      'and for validating all information used to make those decisions.', False, SUB)],
    [('This document is solely for the personal information of the recipient and must not be used '
      'as the basis of any investment decision. Nothing in this document should be construed as '
      'investment or financial advice. The report and information contained herein may not be '
      'altered, reproduced, or redistributed without prior written consent.', False, SUB)],
])

# ── Footer band ──
spacer(doc, 4)
ft = add_table(doc, 1, 1, [W])
no_borders(ft)
fc = ft.cell(0, 0)
shade(fc, PRI)
cell_margins(fc, top=140, bottom=140, left=140, right=140)
write(fc, 'SMC Research — Moneywise. Be Wise.', size=9.5, bold=True, color=WHITE,
      align=WD_ALIGN_PARAGRAPH.CENTER)
write(fc, 'SQE SmallCase · All Indices Portfolio · Monthly Rebalanced', size=8.5,
      color='A9C6E8', align=WD_ALIGN_PARAGRAPH.CENTER, append=True)
write(fc, f'smcresearch.github.io/SQE-/   ·   Generated: {last_update_fmt}', size=8.5,
      color='A9C6E8', align=WD_ALIGN_PARAGRAPH.CENTER, append=True)

doc.save(DOCX_OUT)
print(f'{DOCX_OUT.name} generated — {len(portfolio)} holdings, {len(sector_map)} sectors, '
      f'top-10 weight {top10_weight:.1f}%, updated {last_update_fmt} '
      f'({DOCX_OUT.stat().st_size / 1024:.0f} KB)')
